import hashlib
import json
import os
import uuid
from datetime import datetime
from io import BytesIO  # Для работы с бинарными данными в памяти
from random import randint
from django.template.loader import render_to_string
import numpy as np
import pandas as pd
import psycopg2
from django.conf import settings
from django.contrib import messages
from django.contrib.auth import login, logout
from django.contrib.auth.decorators import login_required
from django.http import FileResponse, JsonResponse
from django.shortcuts import redirect, render
from django.urls import reverse
from django.views import View
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_GET, require_POST
from docx import Document  # Импорт для работы с Word-документами

from .forms import ImprovementProposalForm
from .mixins import AuthRequiredMixin
from .models import *
from .mxl_parser import MxlParser


def get_client_ip(request):
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip


def get_db_connection():
    return psycopg2.connect(
        dbname='tetz',
        user='postgres',
        password='QWE123rty!@#',
        host='localhost',
        port='5432',
        options='-c client_encoding=utf8'
    )

def login_view(request):
    if request.user.is_authenticated:
        return redirect('home')

    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        hashed_password = hashlib.sha256(password.encode('utf-8')).hexdigest()

        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT user_id, username, password FROM users WHERE username = %s", (username,))
            row = cursor.fetchone()
            cursor.close()
            conn.close()

            if row and row[2] == hashed_password:
                try:
                    user = User.objects.get(user_id=row[0])

                    login(request, user)

                    UserActivityLog.objects.create(
                        user=user,
                        action_type='login',
                        ip_address=get_client_ip(request),
                        user_agent=request.META.get('HTTP_USER_AGENT', '')
                    )
                    return redirect('profile')

                except User.DoesNotExist:
                    messages.error(request, 'Пользователь найден в БД, но отсутствует в Django ORM.')
            else:
                messages.error(request, 'Неверное имя пользователя или пароль')

        except Exception as e:
            messages.error(request, f'Ошибка подключения к базе данных: {str(e)}')

    return render(request, 'login.html')


@login_required
def logout_view(request):
    # Логируем выход
    UserActivityLog.objects.create(
        user=request.user,
        action_type='logout',
        ip_address=get_client_ip(request),
        user_agent=request.META.get('HTTP_USER_AGENT', '')
    )
    logout(request)
    messages.info(request, 'Вы успешно вышли из системы')
    return redirect('login')






@login_required
def profile_view(request):
    return render(request, 'profile.html')


def home(request):
    # Путь к Excel файлу с данными
    path = r"C:\Users\kokok\Desktop\diplom_tetz\TETZ\data\production_data.xlsx"

    try:
        # Чтение данных из Excel
        df = pd.read_excel(path, engine='openpyxl')

        # Получаем последние данные (предполагаем определенную структуру Excel)
        monthly_plan = df.iloc[-1]['monthly_plan']  # Последняя запись в столбце monthly_plan
        yearly_plan = df.iloc[-1]['yearly_plan']  # Последняя запись в столбце yearly_plan
        last_update = df.iloc[-1]['date'].strftime('%d.%m.%y')  # Дата последнего обновления

    except Exception as e:
        # Если файл не найден или ошибка чтения - используем значения по умолчанию
        monthly_plan = 100
        yearly_plan = 100
        last_update = "26.12.24"

    context = {
        'monthly_plan': monthly_plan,
        'yearly_plan': yearly_plan,
        'last_update': last_update,
    }
    return render(request, 'home.html', context)

def new_employees(request):
    return render(request, 'new_employees.html', {'title': 'Смотрите, кто пришел'})

def jubilees(request):
    return render(request, 'jubilees.html', {'title': 'Наши Юбиляры'})

def lean_production(request):
    return render(request, 'lean_production.html', {'title': 'Бережливое производство'})


@login_required
def requests_view(request):
    if request.method == 'POST':
        # Обработка формы ППУ
        department = request.POST.get('department')
        problem = request.POST.get('problem_description')
        solution = request.POST.get('proposed_solution')
        expected_result = request.POST.get('expected_result')
        ready_to_implement = request.POST.get('ready_to_implement') == 'on'

        # Создаем запись ППУ
        proposal = ImprovementProposal(
            author=request.user,
            department=department,
            current_situation=problem,
            proposed_solution=solution,
            expected_result=expected_result,
            ready_to_implement=ready_to_implement,
            status='submitted'
        )
        proposal.save()

        # Генерируем документ Word
        doc = Document()
        doc.add_heading('Предложение по улучшению', level=1)

        # Заполняем документ данными
        doc.add_paragraph(f"Зарегистрировано за № {proposal.registration_number}")
        doc.add_paragraph(f"Дата регистрации: {proposal.created_at.strftime('%d.%m.%Y')}")

        doc.add_paragraph("\nРаздел автора:")
        doc.add_paragraph(f"ФИО: {request.user.get_full_name()}")
        doc.add_paragraph(f"Подразделение: {department}")

        doc.add_paragraph("\nОписание проблемы:")
        doc.add_paragraph(problem)

        doc.add_paragraph("\nПредлагаемое решение:")
        doc.add_paragraph(solution)

        doc.add_paragraph("\nОжидаемый результат:")
        doc.add_paragraph(expected_result)

        doc.add_paragraph(f"\nГотов внедрить самостоятельно: {'Да' if ready_to_implement else 'Нет'}")

        # Сохраняем документ
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"ППУ_{proposal.registration_number}.docx"
        proposal.document.save(filename, buffer)
        proposal.save()

        return redirect('requests')

    # Получаем историю заявок пользователя
    proposals = ImprovementProposal.objects.filter(author=request.user).order_by('-created_at')
    return render(request, 'requests.html', {
        'proposals': proposals
    })


@login_required
def download_ppu(request, proposal_id):
    proposal = ImprovementProposal.objects.get(id=proposal_id, author=request.user)
    return FileResponse(proposal.document.open(), as_attachment=True, filename=proposal.document.name)


@require_POST
@login_required
def upload_file(request):
    if not request.FILES.get('file'):
        return JsonResponse({'status': 'error', 'message': 'No file provided'}, status=400)

    try:
        uploaded_file = request.FILES['file']
        file_type = uploaded_file.name.split('.')[-1].lower()

        # Проверяем допустимые форматы файлов
        allowed_types = ['csv', 'xlsx', 'json', 'xml', 'mxl']
        if file_type not in allowed_types:
            return JsonResponse({
                'status': 'error',
                'message': f'Invalid file type. Allowed types: {", ".join(allowed_types)}'
            }, status=400)

        # Создаем папку uploads, если ее нет
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'uploads')
        os.makedirs(upload_dir, exist_ok=True)

        # Генерируем уникальное имя для файла
        unique_filename = f"{uuid.uuid4().hex}.{file_type}"
        file_path = os.path.join(upload_dir, unique_filename)

        # Сохраняем файл
        with open(file_path, 'wb+') as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        # Создаем запись в базе данных
        uploaded_file_record = UploadedFile.objects.create(
            user=request.user,
            original_filename=uploaded_file.name,
            stored_filename=unique_filename,
            file_type=file_type,
            file_size=uploaded_file.size,
            description=request.POST.get('description', '')
        )

        # Логируем действие
        UserActivityLog.objects.create(
            user=request.user,
            action_type='file_upload',
            action_details=f"Uploaded file: {uploaded_file.name}",
            ip_address=get_client_ip(request),
            user_agent=request.META.get('HTTP_USER_AGENT', '')
        )

        # Создаем запись в истории запросов
        QueryHistory.objects.create(
            user=request.user,
            query_text=f"File upload: {uploaded_file.name}",
            query_type='upload',
            file=uploaded_file_record,
            parameters={
                'file_type': file_type,
                'file_size': uploaded_file.size,
                'description': request.POST.get('description', '')
            }
        )

        return redirect(reverse('visualization'))

    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Error uploading file: {str(e)}'
        }, status=500)


@require_POST
@login_required
def delete_file(request, file_id):
    try:
        file_record = UploadedFile.objects.get(file_id=file_id, user=request.user)

        # Удаление файла и записи из БД
        file_path = os.path.join(settings.MEDIA_ROOT, 'uploads', file_record.stored_filename)
        if os.path.exists(file_path):
            os.remove(file_path)
        file_record.delete()

        # Логирование
        UserActivityLog.objects.create(
            user=request.user,
            action_type='file_delete',
            action_details=f"Deleted file: {file_record.original_filename}",
            ip_address=get_client_ip(request),
            user_agent=request.META.get('HTTP_USER_AGENT', '')
        )


        return redirect(reverse('visualization'))

    except Exception as e:
        return JsonResponse(
            {'status': 'error', 'message': str(e)},
            status=400,
            content_type='application/json'
        )

@login_required
def download_file(request, file_id):
    try:
        file_record = UploadedFile.objects.get(file_id=file_id, user=request.user)
        file_path = os.path.join(settings.MEDIA_ROOT, 'uploads', file_record.stored_filename)

        if os.path.exists(file_path):
            # Логируем действие
            UserActivityLog.objects.create(
                user=request.user,
                action_type='file_download',
                action_details=f"Downloaded file: {file_record.original_filename}",
                ip_address=get_client_ip(request),
                user_agent=request.META.get('HTTP_USER_AGENT', '')
            )

            return FileResponse(open(file_path, 'rb'), as_attachment=True, filename=file_record.original_filename)
        else:
            return JsonResponse({'status': 'error', 'message': 'File not found'}, status=404)
    except UploadedFile.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'File not found or access denied'}, status=404)




class VisualizationView(AuthRequiredMixin, View):
    def get(self, request, *args, **kwargs):
        # Получаем список загруженных файлов пользователя
        user_files = UploadedFile.objects.filter(user=request.user).order_by('-upload_date')

        # Получаем сохраненные визуализации пользователя
        user_visualizations = Visualization.objects.filter(user=request.user).order_by('-created_at')

        # Получаем сохраненные конфигурации
        user_configs = VisualizationConfig.objects.filter(user=request.user).order_by('-is_default', 'name')

        context = {
            'title': 'Визуализация данных',
            'user_files': user_files,
            'user_visualizations': user_visualizations,
            'user_configs': user_configs
        }
        return render(request, 'visualization.html', context)


# views.py
@login_required
def visualization_data(request, file_id):
    try:
        file_record = UploadedFile.objects.get(file_id=file_id, user=request.user)
        file_path = os.path.join(settings.MEDIA_ROOT, 'uploads', file_record.stored_filename)

        # Обработка разных типов файлов
        if file_record.file_type == 'mxl':
            return process_mxl_file(request, file_path, file_record)
        elif file_record.file_type == 'xlsx':
            df = pd.read_excel(file_path, engine='openpyxl')
            return process_standard_file(request, file_record, df)
        elif file_record.file_type == 'csv':
            df = pd.read_csv(file_path, sep=';', thousands=' ', encoding='utf-8-sig')
            return process_standard_file(request, file_record, df)

        else:
            messages.error(request, 'Неподдерживаемый формат файла')
            return redirect('visualization')

    except Exception as e:
        messages.error(request, f'Ошибка обработки файла: {str(e)}')
        return redirect('visualization')

def process_standard_file(request, file_record, df):
    """Обработка стандартных файлов (CSV, Excel)"""
    # Очистка данных
    for col in df.columns:
        if df[col].dtype == 'object':
            df[col] = df[col].str.replace(r'[\s\xa0]', '', regex=True)
            try:
                df[col] = pd.to_numeric(df[col])
            except:
                pass

    # Проверка данных
    if df.empty:
        messages.error(request, 'Файл не содержит данных')
        return redirect('visualization')

    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    non_numeric_cols = df.select_dtypes(exclude=[np.number]).columns.tolist()

    if not numeric_cols:
        messages.error(request, 'Файл не содержит числовых данных')
        return redirect('visualization')

    # Подготовка данных для графика
    datasets = [{
        'label': col,
        'data': df[col].tolist()
    } for col in numeric_cols]

    label_column = non_numeric_cols[0] if non_numeric_cols else None
    labels = df[label_column].astype(str).tolist() if label_column else df.index.astype(str).tolist()

    context = {
        'file_id': file_record.file_id,
        'file_name': file_record.original_filename,
        'data': df.to_dict('records'),
        'columns': df.columns.tolist(),
        'chart_data': json.dumps({
            'labels': labels,
            'datasets': datasets
        }, ensure_ascii=False),
        'chart_type': 'bar',
        'user_files': UploadedFile.objects.filter(user=request.user).order_by('-upload_date')
    }
    return render(request, 'visualization.html', context)

def process_mxl_file(request, file_path, file_record):
    """Обработка MXL файлов нового формата"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        parser = MxlParser()
        parsed_data = parser.parse(content)

        # Подготовка данных для графика
        chart_data = {
            'labels': [item['component'] for item in parsed_data['components']],
            'datasets': [
                {
                    'label': 'Укомплектован (%)',
                    'data': [float(item['Укомплектован'].replace(',', '.')) for item in parsed_data['components']],
                    'backgroundColor': 'rgba(75, 192, 192, 0.7)'
                },
                {
                    'label': 'Частично укомплектован (%)',
                    'data': [float(item['Частично укомплектован'].replace(',', '.')) for item in parsed_data['components']],
                    'backgroundColor': 'rgba(255, 206, 86, 0.7)'
                },
                {
                    'label': 'Нет в наличие (%)',
                    'data': [float(item['Нет в наличие'].replace(',', '.')) for item in parsed_data['components']],
                    'backgroundColor': 'rgba(255, 99, 132, 0.7)'
                }
            ]
        }

        # Подготовка данных для таблицы
        table_data = []
        for item in parsed_data['components']:
            table_data.append({
                'Компонент': item['component'],
                'Укомплектован (%)': item['Укомплектован'],
                'Частично укомплектован (%)': item['Частично укомплектован'],
                'Нет в наличие (%)': item['Нет в наличие']
            })

        context = {
            'file_id': file_record.file_id,
            'file_name': file_record.original_filename,
            'data': table_data,
            'columns': ['Компонент', 'Укомплектован (%)', 'Частично укомплектован (%)', 'Нет в наличие (%)'],
            'chart_data': json.dumps(chart_data, ensure_ascii=False),
            'chart_type': 'bar',
            'is_mxl': True,
            'user_files': UploadedFile.objects.filter(user=request.user).order_by('-upload_date')
        }

        return render(request, 'visualization.html', context)

    except Exception as e:
        messages.error(request, f'Ошибка обработки MXL файла: {str(e)}')
        return redirect('visualization')


def process_funnel_data(df):
    # Обработка данных воронки продаж
    labels = df[df['Параметр'].notna()]['Параметр'].tolist()
    values = df[df['Значение'].notna()]['Значение'].tolist()

    return {
        'type': 'funnel',
        'labels': labels,
        'values': values,
        'chart_type': 'bar',
        'options': {
            'indexAxis': 'y',
            'plugins': {
                'title': {
                    'display': True,
                    'text': 'Воронка продаж'
                }
            }
        }
    }


def process_inventory_data(df):
    # Обработка данных комплектации
    summary_rows = df[df['Изделие'].notna() & df['Укомплектован (%)'].notna()]
    products = summary_rows['Изделие'].tolist()
    complete = summary_rows['Укомплектован (%)'].tolist()
    not_available = summary_rows['Нет в наличие (%)'].tolist()

    return {
        'type': 'inventory',
        'products': products,
        'complete': complete,
        'not_available': not_available,
        'chart_type': 'pie',
        'options': {
            'plugins': {
                'title': {
                    'display': True,
                    'text': 'Укомплектованность изделий'
                }
            }
        }
    }


def process_production_data(df):
    # Обработка производственных данных
    months = df['Месяц'].tolist()
    metrics = [col for col in df.columns if col != 'Месяц']
    datasets = []

    for metric in metrics:
        datasets.append({
            'label': metric,
            'data': df[metric].tolist()
        })

    return {
        'type': 'production',
        'labels': months,
        'datasets': datasets,
        'chart_type': 'line',
        'options': {
            'plugins': {
                'title': {
                    'display': True,
                    'text': 'Производственные показатели'
                }
            }
        }
    }


@login_required
@require_POST
def save_visualization(request):
    try:
        # Проверка CSRF
        if not request.POST.get('csrfmiddlewaretoken'):
            return JsonResponse({'status': 'error', 'message': 'CSRF token missing'}, status=400)

        # Проверка обязательных полей
        required_fields = ['title', 'config', 'data', 'file_id']
        for field in required_fields:
            if field not in request.POST:
                return JsonResponse({'status': 'error', 'message': f'Missing required field: {field}'}, status=400)

        try:
            config = json.loads(request.POST['config'])
            data = json.loads(request.POST['data'])
            file_id = int(request.POST['file_id'])
        except (json.JSONDecodeError, ValueError) as e:
            return JsonResponse({'status': 'error', 'message': f'Invalid data format: {str(e)}'}, status=400)

        # Получаем файл и метаданные
        try:
            file_record = UploadedFile.objects.get(file_id=file_id, user=request.user)
            metadata, created = FileDataMetadata.objects.get_or_create(
                file=file_record,
                defaults={
                    'data_type': 'other',
                    'columns_info': {'columns': config.get('tableData', {}).get('columns', [])}
                }
            )
        except UploadedFile.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'File not found'}, status=404)

        # Создаем визуализацию
        visualization = Visualization.objects.create(
            user=request.user,
            metadata=metadata,
            title=request.POST['title'],
            description=request.POST.get('description', ''),
            visualization_type=config.get('chartType', 'bar'),
            config=config,
            is_public=request.POST.get('is_public', 'false') == 'true'
        )

        # Создаем запись в истории
        QueryHistory.objects.create(
            user=request.user,
            query_text=f"Saved visualization: {request.POST['title']}",
            query_type='visualize',
            file=file_record,
            visualization=visualization,
            parameters={
                'chart_type': config.get('chartType'),
                'is_public': request.POST.get('is_public') == 'true'
            }
        )

        return redirect(reverse('visualization'))

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({
            'status': 'error',
            'message': f'Server error: {str(e)}'
        }, status=500)


from django.shortcuts import get_object_or_404


@login_required
def load_visualization(request, viz_id):
    viz = get_object_or_404(Visualization, visualization_id=viz_id, user=request.user)

    table_data = viz.config.get('tableData', {})
    rows = table_data.get('rows', [])
    columns = table_data.get('columns', [])

    structured_rows = []
    if rows and columns:
        # Assuming rows is a flattened list, we need to group by row
        num_columns = len(columns)
        for i in range(0, len(rows), num_columns):
            row_values = rows[i:i + num_columns]
            row_dict = {columns[j]: row_values[j] for j in range(num_columns)}
            structured_rows.append(row_dict)

    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({
            'status': 'success',
            'title': viz.title,
            'chart_data': viz.config['chartData'],
            'chart_type': viz.config.get('chartType', 'bar'),
            'table_data': {
                'columns': columns,
                'rows': structured_rows
            }
        })
    else:
        context = {
            'file_name': viz.title,
            'data': structured_rows,
            'columns': columns,
            'chart_data': json.dumps(viz.config['chartData']),
            'chart_type': viz.config.get('chartType', 'bar'),
            'user_files': UploadedFile.objects.filter(user=request.user),
            'user_visualizations': Visualization.objects.filter(user=request.user)
        }
        return render(request, 'visualization.html', context)


@login_required
@require_POST
def delete_visualization(request, viz_id):
    try:
        visualization = Visualization.objects.get(visualization_id=viz_id, user=request.user)
        visualization.delete()

        return redirect(reverse('visualization'))

    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=400)

# Вспомогательные функции для обработки данных
def process_sales_data(df, chart_type):
    # Анализ данных о продажах
    numeric_cols = df.select_dtypes(include=[np.number]).columns

    if len(numeric_cols) == 0:
        raise ValueError("No numeric columns found for sales data")

    # Если есть дата, группируем по дате
    date_column = None
    for col in df.columns:
        if 'date' in col.lower():
            date_column = col
            break

    if date_column:
        df_grouped = df.groupby(date_column)[numeric_cols].sum().reset_index()
        labels = df_grouped[date_column].astype(str).tolist()
    else:
        df_grouped = df[numeric_cols].sum().to_frame().T
        labels = ['Total']

    datasets = []
    for col in numeric_cols:
        datasets.append({
            'label': col,
            'data': df_grouped[col].tolist(),
            'backgroundColor': get_random_color(),
            'borderColor': get_random_color(),
            'borderWidth': 1
        })

    # Подготовка данных для таблицы
    table_data = []
    for idx, row in df_grouped.iterrows():
        for col in numeric_cols:
            table_data.append({
                'date': str(row[date_column]) if date_column else 'N/A',
                'metric': col,
                'value': row[col],
                'change': np.random.randint(-20, 20)  # Примерное изменение
            })

    return {
        'labels': labels,
        'datasets': datasets,
        'tableData': table_data
    }


def process_inventory_data(df, chart_type):
    # Анализ данных о запасах
    numeric_cols = df.select_dtypes(include=[np.number]).columns

    if len(numeric_cols) == 0:
        raise ValueError("No numeric columns found for inventory data")

    # Для инвентаризации часто используют категории
    category_column = None
    for col in df.columns:
        if 'category' in col.lower() or 'type' in col.lower():
            category_column = col
            break

    if category_column:
        df_grouped = df.groupby(category_column)[numeric_cols].sum().reset_index()
        labels = df_grouped[category_column].astype(str).tolist()
    else:
        labels = numeric_cols.tolist()
        df_grouped = df[numeric_cols].sum().to_frame().T

    datasets = []
    if chart_type == 'pie_chart':
        # Для круговой диаграммы используем первый числовой столбец
        col = numeric_cols[0]
        datasets.append({
            'label': 'Inventory',
            'data': df_grouped[col].tolist(),
            'backgroundColor': [get_random_color() for _ in range(len(labels))]
        })
    else:
        for col in numeric_cols:
            datasets.append({
                'label': col,
                'data': df_grouped[col].tolist(),
                'backgroundColor': get_random_color(),
                'borderColor': get_random_color(),
                'borderWidth': 1
            })

    # Подготовка данных для таблицы
    table_data = []
    for idx, row in df_grouped.iterrows():
        for col in numeric_cols:
            table_data.append({
                'date': str(row[category_column]) if category_column else 'N/A',
                'metric': col,
                'value': row[col],
                'change': np.random.randint(-10, 10)  # Примерное изменение
            })

    return {
        'labels': labels,
        'datasets': datasets,
        'tableData': table_data
    }


def process_customers_data(df, chart_type):
    # Анализ данных о клиентах
    categorical_cols = df.select_dtypes(include=['object', 'category']).columns

    if len(categorical_cols) == 0:
        raise ValueError("No categorical columns found for customers data")

    # Выбираем первый категориальный столбец для группировки
    group_col = categorical_cols[0]
    value_cols = [col for col in df.columns if col != group_col and df[col].dtype in [np.number, 'int', 'float']]

    if not value_cols:
        # Если нет числовых столбцов, считаем количество
        df_grouped = df[group_col].value_counts().reset_index()
        df_grouped.columns = [group_col, 'count']
        value_cols = ['count']
    else:
        df_grouped = df.groupby(group_col)[value_cols].sum().reset_index()

    labels = df_grouped[group_col].astype(str).tolist()

    datasets = []
    for col in value_cols:
        datasets.append({
            'label': col,
            'data': df_grouped[col].tolist(),
            'backgroundColor': get_random_color(),
            'borderColor': get_random_color(),
            'borderWidth': 1
        })

    # Подготовка данных для таблицы
    table_data = []
    for idx, row in df_grouped.iterrows():
        for col in value_cols:
            table_data.append({
                'date': str(row[group_col]),
                'metric': col,
                'value': row[col],
                'change': np.random.randint(-15, 15)  # Примерное изменение
            })

    return {
        'labels': labels,
        'datasets': datasets,
        'tableData': table_data
    }


def process_generic_data(df, chart_type):
    # Общая обработка для неизвестных типов данных
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    categorical_cols = df.select_dtypes(include=['object', 'category']).columns

    if len(numeric_cols) == 0:
        # Если нет числовых столбцов, используем категориальные
        if len(categorical_cols) == 0:
            raise ValueError("No numeric or categorical columns found")

        group_col = categorical_cols[0]
        df_grouped = df[group_col].value_counts().reset_index()
        df_grouped.columns = [group_col, 'count']
        labels = df_grouped[group_col].astype(str).tolist()

        datasets = [{
            'label': 'Count',
            'data': df_grouped['count'].tolist(),
            'backgroundColor': [get_random_color() for _ in labels],
            'borderColor': [get_random_color() for _ in labels],
            'borderWidth': 1
        }]
    else:
        # Используем числовые столбцы
        labels = df.index.astype(str).tolist() if len(df) < 20 else []
        datasets = []
        for col in numeric_cols:
            datasets.append({
                'label': col,
                'data': df[col].tolist(),
                'backgroundColor': get_random_color(),
                'borderColor': get_random_color(),
                'borderWidth': 1
            })

    # Подготовка данных для таблицы
    table_data = []
    sample_col = numeric_cols[0] if len(numeric_cols) > 0 else 'count'
    for idx, row in df.iterrows():
        table_data.append({
            'date': str(idx),
            'metric': sample_col,
            'value': row[sample_col] if len(numeric_cols) > 0 else idx,
            'change': np.random.randint(-5, 5)
        })

    return {
        'labels': labels,
        'datasets': datasets,
        'tableData': table_data[:100]  # Ограничиваем количество строк для таблицы
    }


def process_json_data(json_data, chart_type):
    # Обработка JSON данных (упрощенная версия)
    if isinstance(json_data, list):
        # Если JSON представляет собой массив объектов
        df = pd.DataFrame(json_data)
        return process_generic_data(df, chart_type)
    elif isinstance(json_data, dict):
        # Если JSON представляет собой объект
        labels = list(json_data.keys())
        datasets = [{
            'label': 'Values',
            'data': list(json_data.values()),
            'backgroundColor': [get_random_color() for _ in labels],
            'borderColor': [get_random_color() for _ in labels],
            'borderWidth': 1
        }]

        table_data = []
        for key, value in json_data.items():
            table_data.append({
                'date': key,
                'metric': 'value',
                'value': value,
                'change': np.random.randint(-10, 10)
            })

        return {
            'labels': labels,
            'datasets': datasets,
            'tableData': table_data
        }
    else:
        raise ValueError("Unsupported JSON structure")


def get_random_color():
    # Генерация случайного цвета
    return f'rgba({np.random.randint(0, 255)}, {np.random.randint(0, 255)}, {np.random.randint(0, 255)}, 0.7)'


@login_required
def improvement_proposals(request):
    if request.method == 'POST':
        form = ImprovementProposalForm(request.POST, request.FILES)
        if form.is_valid():
            proposal = form.save(commit=False)
            proposal.author = request.user
            proposal.save()
            return redirect('generate_ppu_document', proposal_id=proposal.id)
    else:
        form = ImprovementProposalForm()

    proposals = ImprovementProposal.objects.filter(author=request.user).order_by('-created_at')
    return render(request, 'improvement_proposals.html', {
        'form': form,
        'proposals': proposals
    })


@login_required
def generate_ppu_document(request, proposal_id):
    proposal = ImprovementProposal.objects.get(id=proposal_id, author=request.user)

    # Создаем документ
    doc = Document()

    # Заголовок
    title = doc.add_heading('Предложение по улучшению', level=1)
    title.alignment = 1  # CENTER

    # Регистрационные данные
    doc.add_paragraph(f"Зарегистрировано за № {proposal.registration_number}")
    doc.add_paragraph(
        f"Дата регистрации «{proposal.created_at.day}» {proposal.created_at.strftime('%B')} {proposal.created_at.year} г.")

    # Раздел автора
    author_section = doc.add_paragraph()
    author_section.add_run("Раздел автора\n").bold = True
    author_section.add_run(
        f"Дата подачи предложения «{proposal.submission_date.day}» {proposal.submission_date.strftime('%B')} {proposal.submission_date.year} г.\n")

    author_info = [
        f"Фамилия, имя, отчество автора: {proposal.author.get_full_name()}",
        f"Должность: {proposal.author.position if hasattr(proposal.author, 'position') else 'Не указана'}",
        f"Подразделение: {proposal.department}",
        f"Коэффициент участия: {proposal.participation_coefficient}%"
    ]

    for info in author_info:
        doc.add_paragraph(info)

    # Содержание предложения
    doc.add_heading('ИДЕЯ АВТОРА', level=2)
    doc.add_paragraph(proposal.current_situation)

    if proposal.current_situation_sketch:
        doc.add_paragraph("Эскиз текущей ситуации прилагается")

    doc.add_heading('Предлагаемое решение:', level=2)
    doc.add_paragraph(proposal.proposed_solution)

    if proposal.solution_sketch:
        doc.add_paragraph("Эскиз решения прилагается")

    doc.add_heading('Ожидаемый результат (эффект):', level=2)
    doc.add_paragraph(proposal.expected_result)

    # Заключительная часть
    implementation = doc.add_paragraph()
    implementation.add_run("Готов ли автор (авторы) самостоятельно внедрить решение: ").bold = True
    implementation.add_run("Да" if proposal.ready_to_implement else "Нет")

    doc.add_paragraph(f"Необходимые ресурсы: {proposal.required_resources}")

    # Сохраняем документ
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Обновляем запись
    filename = f"ППУ_{proposal.registration_number}_{datetime.now().strftime('%Y%m%d')}.docx"
    proposal.document.save(filename, buffer)
    proposal.status = 'submitted'
    proposal.save()

    return FileResponse(buffer, as_attachment=True, filename=filename)


def process_1c_mxl(file_path):
    """Обработка MXL-файлов комплектации из 1С"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        parser = MxlParser()
        data = parser.parse(content)

        # Подготовка данных для визуализации
        labels = [item['name'] for item in data['items']]
        required = [int(item['required'].replace('\u00a0', '')) for item in data['items']]
        available = [int(item['available'].replace('\u00a0', '')) for item in data['items']]

        return {
            'status': 'success',
            'data_type': '1c_complectation',
            'header': data['header'],
            'summary': data['summary'],
            'labels': labels,
            'datasets': [
                {
                    'label': 'Требуется',
                    'data': required,
                    'backgroundColor': 'rgba(255, 99, 132, 0.7)'
                },
                {
                    'label': 'Доступно',
                    'data': available,
                    'backgroundColor': 'rgba(54, 162, 235, 0.7)'
                }
            ],
            'tableData': data['items']
        }

    except Exception as e:
        return {'status': 'error', 'message': str(e)}